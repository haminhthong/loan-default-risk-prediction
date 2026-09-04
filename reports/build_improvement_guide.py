"""Tạo báo cáo DOCX hướng dẫn cải thiện dự án theo bốn tầng."""

from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "Loan_Default_Risk_Improvement_Guide.docx"

NAVY = RGBColor(31, 78, 121)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(89, 96, 105)
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F2F4F7"
GREEN = "E7F3EA"
AMBER = "FFF4CE"
RED = "FDE9E7"


def set_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_font(run, 8.5)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        shade(cell, "2E74B5")
        for run in cell.paragraphs[0].runs:
            set_font(run, 8.7, bold=True, color=RGBColor(255, 255, 255))
    for i, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if i % 2:
            for cell in row.cells:
                shade(cell, LIGHT_GRAY)
        for cell, value in zip(row.cells, row_data):
            cell.text = str(value)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}: ")
    set_font(r, 10, bold=True, color=NAVY)
    r = p.add_run(text)
    set_font(r, 10)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, "F6F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for line in text.strip().splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
        run.font.size = Pt(8.3)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_font(run, 8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for style_name, size, before, after, color in [
        ("Heading 1", 16, 14, 6, BLUE),
        ("Heading 2", 13, 10, 4, BLUE),
        ("Heading 3", 11.5, 8, 3, NAVY),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    header.text = "LOAN DEFAULT RISK  |  HƯỚNG DẪN CẢI THIỆN DỰ ÁN"
    for run in header.runs:
        set_font(run, 8.2, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])


def build():
    doc = Document()
    configure_document(doc)

    # Editorial cover.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HƯỚNG DẪN CẢI THIỆN CHI TIẾT")
    set_font(r, 13, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dự Án Dự Báo Rủi Ro\nVỡ Nợ Khoản Vay")
    set_font(r, 28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Problem  →  AI/ML Correctness  →  Software Engineering  →  Production & Business Value")
    set_font(r, 11.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    add_callout(
        doc,
        "Mục tiêu",
        "Biến một dự án portfolio có nền tảng ML tốt thành repository tái lập được, có kiểm thử đáng tin cậy và có bằng chứng vận hành; không phóng đại thành hệ thống tín dụng production.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Phiên bản: 01/09/2026  |  Phạm vi: Portfolio / CV / Phỏng vấn kỹ thuật")
    set_font(r, 9.5, color=GRAY)
    doc.add_page_break()

    doc.add_heading("1. Kết luận điều hành", level=1)
    doc.add_paragraph(
        "Dự án đã giải quyết đúng bài toán phân loại Charged Off theo thời điểm cấp vay, có temporal holdout, baseline, calibration và train-serving parity. Tuy nhiên, dự án chưa đủ điều kiện để tuyên bố production-ready vì người khác chưa thể clone rồi tái lập ngay, test edge case còn mỏng, nguồn dữ liệu chưa xác minh và chưa có load test/security controls."
    )
    add_table(
        doc,
        ["Tầng", "Hiện trạng", "Ưu tiên quan trọng nhất"],
        [
            ("Problem", "Tốt cho portfolio", "Chuẩn hóa problem statement và data contract"),
            ("AI/ML", "Đúng nền tảng", "Temporal CV, leakage guard hai lớp, CI metric"),
            ("Software", "Có cấu trúc", "Clone-to-run, edge tests, API contract"),
            ("Production", "Chưa chứng minh", "Load test, security, monitoring, governance"),
        ],
        [1600, 2400, 5360],
    )
    add_callout(doc, "Nguyên tắc ưu tiên", "Không thêm XGBoost, SHAP hoặc giao diện mới trước khi hoàn thành tính tái lập, kiểm thử và data provenance.", AMBER)

    doc.add_heading("2. Tầng 1 — Problem", level=1)
    doc.add_heading("2.1. Chuẩn hóa định nghĩa bài toán", level=2)
    doc.add_paragraph("README và model card cần thống nhất sáu thành phần sau:")
    add_table(
        doc,
        ["Thành phần", "Định nghĩa nên dùng"],
        [
            ("Mục tiêu", "Ước lượng xác suất một khoản vay trở thành Charged Off."),
            ("Đơn vị dự báo", "Một khoản vay, không phải một khách hàng."),
            ("Thời điểm dự báo", "Ngay trước hoặc tại thời điểm phát hành khoản vay."),
            ("Nhãn", "Charged Off = 1; Fully Paid = 0; Current bị loại."),
            ("Đầu ra", "Xác suất rủi ro và cờ cảnh báo theo threshold validation."),
            ("Ngoài phạm vi", "Không tự động duyệt/từ chối, định giá hoặc thay chuyên viên."),
        ],
        [2200, 7160],
    )
    add_callout(doc, "Cách diễn đạt", "Dùng “Charged Off risk” hoặc “rủi ro khoản vay bị ghi giảm” khi cần chính xác; không khẳng định đây là định nghĩa default pháp lý hoặc Basel.")

    doc.add_heading("2.2. Data contract cho input", level=2)
    add_table(
        doc,
        ["Trường", "Kiểu/đơn vị", "Thời điểm có", "Validation tối thiểu"],
        [
            ("loan_amnt", "float, USD", "Khi đăng ký", "> 0; giới hạn cực trị"),
            ("annual_inc", "float, USD/năm", "Khi đăng ký", "> 0; không inf"),
            ("dti", "float, %", "Khi đăng ký", "0–100 hoặc theo data card"),
            ("term", "category, tháng", "Khi đăng ký", "36 hoặc 60 months"),
            ("grade", "category", "Khi thẩm định", "A–G"),
            ("issue_d", "MMM-YY", "Khi phát hành", "Parse được và hợp lệ"),
        ],
        [1600, 1900, 2400, 3460],
    )
    code_block(doc, '''from typing import Literal

term: Literal["36 months", "60 months"]
grade: Literal["A", "B", "C", "D", "E", "F", "G"]
addr_state: str = Field(pattern=r"^[A-Z]{2}$")
dti: float | None = Field(default=None, ge=0, le=100)''')

    doc.add_heading("2.3. Kiểm tra policy proxy", level=2)
    doc.add_paragraph("Bỏ int_rate và sub_grade là bước tốt nhưng chưa loại hết dấu vết chính sách underwriting cũ. Cần chạy ablation theo bốn cấu hình:")
    for text in [
        "Tất cả đặc trưng hiện tại.",
        "Không có int_rate và sub_grade.",
        "Không có int_rate, sub_grade và grade.",
        "Không có int_rate, sub_grade, grade và installment.",
    ]:
        add_number(doc, text)
    doc.add_paragraph("So sánh PR-AUC, ROC-AUC, recall, precision, Brier score và drift. Nếu hiệu năng giảm mạnh khi bỏ grade, phải ghi rõ model phụ thuộc vào hệ thống xếp hạng cũ.")

    doc.add_heading("3. Tầng 2 — AI/ML Correctness", level=1)
    doc.add_heading("3.1. Leakage guard hai lớp", level=2)
    doc.add_paragraph("Giữ allowlist hiện tại và bổ sung denylist để một cột hậu nghiệm không thể lọt vào model do thay đổi schema.")
    code_block(doc, '''LEAKAGE_COLUMNS = {
    "loan_status", "default_flag", "total_pymnt", "recoveries",
    "collection_recovery_fee", "last_pymnt_d", "last_pymnt_amnt",
    "next_pymnt_d", "out_prncp", "out_prncp_inv",
}

unexpected = LEAKAGE_COLUMNS.intersection(features.columns)
if unexpected:
    raise ValueError(f"Phát hiện cột hậu nghiệm: {sorted(unexpected)}")''')
    add_bullet(doc, "Test bắt buộc: target và toàn bộ denylist không xuất hiện sau build_features.")
    add_bullet(doc, "Kiểm tra feature availability bằng data contract, không chỉ dựa vào tên cột.")
    add_bullet(doc, "Phân biệt target leakage với policy proxy; grade không phải hậu nghiệm nhưng có thể học lại chính sách cũ.")

    doc.add_heading("3.2. Temporal cross-validation", level=2)
    doc.add_paragraph("Test out-of-time hiện tại đúng, nhưng model comparison vẫn dùng stratified random CV trong tập train. Nên chuyển sang expanding-window folds theo tháng:")
    add_table(
        doc,
        ["Fold", "Train", "Validation nội bộ"],
        [
            ("1", "2007–2008", "2009"),
            ("2", "2007–2009", "01–06/2010"),
            ("3", "2007–06/2010", "07–12/2010"),
        ],
        [1200, 3880, 4280],
    )
    add_callout(doc, "Quy tắc", "Không để các quan sát cùng tháng xuất hiện ở cả train và validation. Giữ H1/2011 để chọn threshold và H2/2011 làm test cuối.")

    doc.add_heading("3.3. Target censoring", level=2)
    doc.add_paragraph("Loại Current hợp lý cho supervised learning, nhưng có thể tạo selection bias. Cần báo cáo theo issue_month: tổng khoản vay, Fully Paid, Charged Off, Current và tỷ lệ bị loại. Nếu tháng cuối có tỷ lệ Current cao, test chỉ đại diện cho nhóm khoản vay đã kết thúc quan sát.")

    doc.add_heading("3.4. Threshold theo chi phí", level=2)
    doc.add_paragraph("Công thức FN × 5 + FP × 1 hợp lệ như sensitivity scenario, nhưng chưa phải Expected Loss. Cách cải thiện:")
    add_bullet(doc, "Giữ bảng độ nhạy 2:1, 5:1, 10:1 và gọi threshold 0,14 là ngưỡng minh họa.")
    add_bullet(doc, "Khi có dữ liệu nghiệp vụ, thay FN cost bằng EAD × LGD và FP cost bằng lợi nhuận cơ hội bị mất.")
    add_bullet(doc, "Bổ sung capacity constraint: số hồ sơ cảnh báo không vượt năng lực manual review.")
    code_block(doc, '''Expected loss bỏ sót = Probability of Default × EAD × LGD
Net value(threshold) = Loss avoided - Review cost - Opportunity cost''')

    doc.add_heading("3.5. Độ bất định và calibration", level=2)
    add_bullet(doc, "Bootstrap 1.000 lần để báo cáo 95% CI cho PR-AUC, ROC-AUC, recall và precision.")
    add_bullet(doc, "Không kết luận 0,3384 tốt hơn chắc chắn 0,3366 nếu confidence intervals chồng lấn.")
    add_bullet(doc, "Bổ sung calibration curve, ECE, calibration slope/intercept và Brier skill score.")
    add_bullet(doc, "Gắn cảnh báo vùng xác suất 0,2–0,3 đang đánh giá thấp rủi ro trên test.")

    doc.add_heading("3.6. Fairness đúng phạm vi", level=2)
    doc.add_paragraph("Đổi tên slice report thành “Performance stability by segment”. Không tuyên bố fairness audit khi thiếu thuộc tính nhạy cảm, cơ sở pháp lý và quy trình governance.")

    doc.add_heading("4. Tầng 3 — Software Engineering", level=1)
    doc.add_heading("4.1. Clone-to-run và xác minh dữ liệu", level=2)
    doc.add_paragraph("Đây là khoảng trống quan trọng nhất. Data và artifact bị loại khỏi Git nên người clone chưa thể chạy ngay. Thêm scripts/verify_data.py để kiểm tra:")
    for text in ["File tồn tại và SHA-256 đúng.", "Số dòng/cột và cột bắt buộc đúng.", "Khoảng issue_d và phân phối loan_status đúng.", "Thông báo rõ cách tự đặt dữ liệu khi không được phân phối lại."]:
        add_bullet(doc, text)
    code_block(doc, '''git clone <repository-url>
cd loan-default-risk-prediction
cp lendingclub_2007_2011.csv data/raw/
python scripts/verify_data.py
python -m src.train
python -m pytest -q''')
    add_callout(doc, "Không được làm", "Không tạo URL tải dataset giả hoặc công khai CSV khi giấy phép chưa được xác minh.", RED)

    doc.add_heading("4.2. Mở rộng test suite", level=2)
    add_table(
        doc,
        ["Nhóm", "Edge cases cần thêm"],
        [
            ("Data", "Thiếu cột; ID rỗng; status lạ; ngày sai; split rỗng"),
            ("Features", "% sai; income = 0; ngày tín dụng tương lai; leakage denylist"),
            ("Threshold", "Input rỗng; probability ngoài [0,1]; một lớp; tie deterministic"),
            ("Inference", "Unknown category; optional missing; artifact thiếu key; parity"),
            ("API", "Valid score; empty batch; >1.000 records; model missing; invalid enum"),
            ("UI", "File quá lớn; CSV hỏng; thiếu schema; số dòng vượt giới hạn"),
        ],
        [1800, 7560],
    )
    add_callout(doc, "Mục tiêu", "Khoảng 25–35 test có ý nghĩa. Không chạy theo coverage 100% nếu test chỉ lặp lại implementation.")

    doc.add_heading("4.3. API contract và lỗi", level=2)
    add_bullet(doc, "Đổi ConfigDict(extra='ignore') thành extra='forbid' để bắt lỗi gõ sai tên trường.")
    add_bullet(doc, "Dùng enum/Literal cho term, grade, home_ownership và verification_status.")
    add_bullet(doc, "Không trả model_path tuyệt đối trong /health; chỉ trả model_loaded và version.")
    add_bullet(doc, "Kiểm tra artifact có đủ pipeline, threshold, feature_columns và model_name trước khi dùng.")
    add_bullet(doc, "Không load joblib do người dùng upload vì deserialization có thể thực thi mã không tin cậy.")
    code_block(doc, '''return {
    "status": "ok",
    "model_loaded": MODEL_PATH.exists(),
    "model_version": artifact.get("model_version"),
}''')

    doc.add_heading("4.4. CSV upload và CI", level=2)
    add_bullet(doc, "Giới hạn file 10 MB và tối đa 10.000 dòng cho demo; kiểm tra schema trước inference.")
    add_bullet(doc, "Không hiển thị stack trace hoặc đường dẫn nội bộ cho người dùng.")
    add_bullet(doc, "CI chạy pytest, compileall, Docker build và smoke test /health.")
    add_bullet(doc, "Artifact fixture nhỏ nên được dùng cho integration test; artifact thật có thể ở release/model registry.")

    doc.add_heading("5. Tầng 4 — Production & Business Value", level=1)
    doc.add_heading("5.1. Chứng minh khả năng chịu 100 users", level=2)
    doc.add_paragraph("Không thể suy ra khả năng tải từ việc dự án dùng FastAPI và Docker. Thêm Locust hoặc k6 và đặt acceptance criteria trước khi chạy:")
    add_table(
        doc,
        ["Chỉ số", "Mục tiêu ban đầu"],
        [
            ("Concurrent users", "100"),
            ("p95 latency", "< 500 ms cho batch 1 hồ sơ"),
            ("Error rate", "< 1%"),
            ("Throughput", ">= 50 requests/giây, tùy máy benchmark"),
            ("Tài nguyên", "RAM ổn định; CPU không bão hòa kéo dài"),
        ],
        [2600, 6760],
    )
    code_block(doc, '''locust -f load_tests/locustfile.py --headless \\
  --users 100 --spawn-rate 10 --run-time 2m \\
  --host http://localhost:8000''')
    add_bullet(doc, "Benchmark 1, 2 và 4 worker; không giả định càng nhiều worker càng tốt.")
    add_bullet(doc, "Đo riêng batch 1, batch 10 và batch 1.000 hồ sơ.")

    doc.add_heading("5.2. Security và privacy", level=2)
    add_table(
        doc,
        ["Rủi ro", "Kiểm soát đề xuất"],
        [
            ("API công khai", "API key/OAuth, rate limit, reverse proxy HTTPS"),
            ("Upload lớn", "Byte limit, row limit, timeout, schema validation"),
            ("Log PII", "Không log body; dùng request ID và log giảm định danh"),
            ("Artifact độc hại", "Chỉ load artifact nội bộ đã ký/checksum"),
            ("Lưu dữ liệu", "Retention policy, mã hóa, xóa file tạm"),
            ("Swagger", "Hạn chế truy cập trong môi trường public"),
        ],
        [2600, 6760],
    )

    doc.add_heading("5.3. Monitoring và model governance", level=2)
    add_bullet(doc, "Runtime: request count, error rate, p50/p95/p99, batch size, CPU và RAM.")
    add_bullet(doc, "Data: missing rate, unknown-category rate, prediction distribution và alert rate.")
    add_bullet(doc, "ML: PSI theo cửa sổ, calibration khi có nhãn và performance theo segment.")
    add_bullet(doc, "Governance: model version, data checksum, git commit, library versions, approval và rollback.")
    code_block(doc, '''artifact_metadata = {
    "model_version": "1.0.0",
    "trained_at_utc": "...",
    "data_sha256": "...",
    "git_commit": "...",
    "python_version": "...",
    "sklearn_version": "1.9.0",
    "feature_schema_version": "1.0",
}''')
    add_callout(doc, "Drift", "PSI cao phải kích hoạt điều tra, không tự động tái huấn luyện hoặc triển khai model mới.")

    doc.add_heading("5.4. Chứng minh giá trị kinh doanh", level=2)
    doc.add_paragraph("Mỗi threshold cần được chuyển thành tác động vận hành:")
    add_bullet(doc, "Số hồ sơ bị cảnh báo và số hồ sơ chuyển manual review.")
    add_bullet(doc, "Năng lực xử lý của chuyên viên và SLA.")
    add_bullet(doc, "Nợ xấu ước tính tránh được, khách hàng tốt bị ảnh hưởng và lợi nhuận ròng.")
    add_bullet(doc, "Ngưỡng tối ưu toán học không được chọn nếu vượt capacity hoặc policy constraints.")

    doc.add_heading("6. Lộ trình triển khai", level=1)
    add_table(
        doc,
        ["Giai đoạn", "Thời lượng gợi ý", "Đầu ra bắt buộc", "Điều kiện hoàn thành"],
        [
            ("1. CV-ready", "2–4 ngày", "Data contract, verify_data, API fixes, 25+ tests", "Clone có hướng dẫn tái lập; CI xanh"),
            ("2. Nổi bật", "3–6 ngày", "Temporal CV, bootstrap CI, ablation, load test", "Metric có CI; báo cáo tải 100 users"),
            ("3. Production-oriented", "1–2 tuần", "Auth, monitoring, privacy, registry", "Có threat model, dashboard và rollback"),
        ],
        [1500, 1800, 3600, 2460],
    )

    doc.add_heading("7. Backlog ưu tiên theo tác động", level=1)
    add_table(
        doc,
        ["P", "Công việc", "Tác động", "Độ khó"],
        [
            ("P0", "verify_data.py + hướng dẫn dữ liệu", "Tái lập và uy tín", "Thấp"),
            ("P0", "API extra=forbid + enum + ẩn path", "Độ tin cậy/security", "Thấp"),
            ("P0", "Leakage denylist + tests", "ML correctness", "Thấp"),
            ("P1", "25–35 edge/integration tests", "Ổn định phần mềm", "Trung bình"),
            ("P1", "Temporal CV + bootstrap CI", "Độ tin cậy metric", "Trung bình"),
            ("P1", "Ablation grade/installment", "Minh bạch policy proxy", "Trung bình"),
            ("P1", "Locust 100 users", "Bằng chứng hiệu năng", "Trung bình"),
            ("P2", "Auth/rate limit/monitoring", "Gần production", "Cao"),
            ("P2", "Expected Loss + capacity", "Giá trị nghiệp vụ", "Cao"),
        ],
        [600, 4000, 3100, 1660],
    )

    doc.add_heading("8. Checklist nghiệm thu", level=1)
    checklist = [
        "Problem statement nêu đúng đơn vị, thời điểm dự báo, nhãn và intended use.",
        "Data contract mô tả kiểu, đơn vị, miền hợp lệ và thời điểm có của mọi input.",
        "Allowlist và denylist đều ngăn cột hậu nghiệm.",
        "Temporal folds không trộn cùng tháng giữa train và validation.",
        "Threshold được gọi đúng là sensitivity scenario nếu chưa có EAD/LGD.",
        "Metric có confidence interval và calibration diagnostics.",
        "Người clone có quy trình verify data → train → test rõ ràng.",
        "API từ chối field thừa, enum sai, batch rỗng và batch quá lớn.",
        "Health endpoint không lộ đường dẫn máy chủ hoặc secret.",
        "Test suite bao phủ edge case và API integration.",
        "Docker được build và smoke-test trong CI.",
        "Load test 100 users có p95, throughput, error rate và tài nguyên.",
        "Không log dữ liệu hồ sơ; có retention và xóa file upload.",
        "Artifact có version, checksum dữ liệu, git commit và library versions.",
        "README không dùng production-ready trước khi có đủ bằng chứng.",
    ]
    for item in checklist:
        add_bullet(doc, "☐ " + item)

    doc.add_heading("9. Cách trình bày sau khi cải thiện", level=1)
    add_callout(
        doc,
        "CV bullet",
        "Xây dựng leakage-safe credit-risk pipeline với temporal validation, calibrated probability và cost-sensitive threshold; đóng gói train-serving parity, automated tests, Docker CI và load-test evidence.",
        GREEN,
    )
    doc.add_paragraph("Trong phỏng vấn, trình bày theo chuỗi: vấn đề → rủi ro leakage → thiết kế thời gian → metric và uncertainty → threshold nghiệp vụ → khả năng tái lập → giới hạn. Giá trị nổi bật nằm ở protocol đánh giá trung thực, không phải số lượng thuật toán.")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
